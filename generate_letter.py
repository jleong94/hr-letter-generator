#!/usr/bin/env python3
"""
HR Letter Generator - Resignation Letter (English / Bahasa Melayu / 中文 简体)

Prompts for the required details, auto-calculates the last working day from the
notice date + notice period (in days, weeks, or months), and writes a professional resignation
letter as both Word (.docx) and PDF (.pdf) into the ./generated_letters folder.
The letter content is generated in the language chosen at the start.

Alongside the letter it also writes two ready-to-send email drafts (.txt), in the
same language:
  1. a short cover email to send WITH the letter attached, and
  2. a full-text email that carries the whole resignation in the body itself,
     so no attachment is needed.

This script is normally launched by run.cmd (Windows) or run.command / run.sh
(macOS), which set up Python and the required packages automatically.
"""

import os
import re
import sys
import calendar
import subprocess
from pathlib import Path
from datetime import date, datetime, timedelta

try:
    from docx import Document
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
except ImportError as exc:  # pragma: no cover - handled by the launchers
    print()
    print("Required packages are missing:", exc)
    print("Please start this tool with run.cmd (Windows) or run.command / run.sh (Mac).")
    print("Those launchers install everything automatically.")
    sys.exit(1)


# --------------------------------------------------------------------------- #
# Language data
# --------------------------------------------------------------------------- #
LANGUAGE_CHOICES = {"1": "en", "2": "ms", "3": "zh", "": "en"}

# Notice-period units. The console UI is English (like every other prompt);
# only the finished letter is translated.
NOTICE_UNIT_CHOICES = {"1": "day", "2": "week", "3": "month", "": "month"}
UNIT_PLURAL_EN = {"day": "days", "week": "weeks", "month": "months"}

TRANSLATIONS = {
    "en": {
        "filename_prefix": "Resignation_Letter",
        "title": "LETTER OF RESIGNATION",
        "subject": "Re: Resignation from the Position of {position}",
        "salutation_named": "Dear {name},",
        "salutation_generic": "Dear Sir/Madam,",
        "opening_notice": (
            "Please accept this letter as formal notification of my resignation from the "
            "position of {position} at {company}. In accordance with the terms of my "
            "employment, I will serve a notice period of {period}, with my last "
            "working day being {last_day}."
        ),
        "opening_immediate": (
            "Please accept this letter as formal notification of my resignation from the "
            "position of {position} at {company}, with immediate effect. My last working "
            "day will be {last_day}."
        ),
        "gratitude_ach": (
            "My time at {company} has been truly rewarding. I am especially grateful for the "
            "opportunity to {achievements}, as well as for the professional growth, "
            "mentorship, and collaboration I have enjoyed throughout my tenure."
        ),
        "gratitude_noach": (
            "My time at {company} has been truly rewarding, and I am grateful for the "
            "professional growth, mentorship, and collaboration I have enjoyed throughout "
            "my tenure."
        ),
        "handover": (
            "I am committed to ensuring a smooth and professional transition during my "
            "remaining time. I am glad to assist with handing over my responsibilities, "
            "documenting ongoing work, training a successor, and completing any outstanding "
            "deliverables to the best of my ability."
        ),
        "closing_para": (
            "Thank you for your guidance and support during my time here. I wish you and "
            "{company} continued success, and I hope to remain in touch."
        ),
        "signoff": "Yours sincerely,",
        # --- email drafts: header labels, subject, and the cover-email body ---
        "email_label_to": "To",
        "email_label_subject": "Subject",
        "email_subject": "Resignation Notice - {your_name} ({position})",
        "email_cover_notice": (
            "I am writing to formally notify you of my resignation from the position of "
            "{position} at {company}. My formal letter of resignation is attached to this "
            "email for your records. In line with my notice period of {period}, my last "
            "working day will be {last_day}."
        ),
        "email_cover_immediate": (
            "I am writing to formally notify you of my resignation from the position of "
            "{position} at {company}, with immediate effect. My formal letter of resignation "
            "is attached to this email for your records, and my last working day will be "
            "{last_day}."
        ),
        "email_cover_body": (
            "I am committed to ensuring a smooth handover before I leave, and I am happy to "
            "assist in any way to make the transition as seamless as possible. Thank you for "
            "the support and guidance I have received during my time at {company}."
        ),
    },
    "ms": {
        "filename_prefix": "Surat_Peletakan_Jawatan",
        "title": "SURAT PELETAKAN JAWATAN",
        "subject": "Perkara: Peletakan Jawatan sebagai {position}",
        "salutation_named": "Yang dihormati {name},",
        "salutation_generic": "Tuan/Puan,",
        "opening_notice": (
            "Dengan segala hormatnya, surat ini merupakan notis rasmi peletakan jawatan saya "
            "sebagai {position} di {company}. Selaras dengan terma pekerjaan saya, saya akan "
            "memenuhi tempoh notis selama {period}, dengan hari terakhir saya bekerja pada "
            "{last_day}."
        ),
        "opening_immediate": (
            "Dengan segala hormatnya, surat ini merupakan notis rasmi peletakan jawatan saya "
            "sebagai {position} di {company}, berkuat kuasa serta-merta. Hari terakhir saya "
            "bekerja ialah {last_day}."
        ),
        "gratitude_ach": (
            "Tempoh perkhidmatan saya di {company} amat bermakna. Saya amat menghargai peluang "
            "untuk {achievements}, serta perkembangan profesional, bimbingan, dan semangat "
            "kerjasama yang saya nikmati sepanjang tempoh ini."
        ),
        "gratitude_noach": (
            "Tempoh perkhidmatan saya di {company} amat bermakna, dan saya menghargai "
            "perkembangan profesional, bimbingan, dan semangat kerjasama yang saya nikmati "
            "sepanjang tempoh ini."
        ),
        "handover": (
            "Saya komited untuk memastikan proses peralihan tugas berjalan lancar dan "
            "profesional sepanjang baki tempoh perkhidmatan saya. Saya sedia membantu "
            "menyerahkan tanggungjawab, mendokumentasikan kerja yang sedang berjalan, melatih "
            "pengganti, serta menyiapkan sebarang tugasan tertunggak dengan sebaik mungkin."
        ),
        "closing_para": (
            "Terima kasih atas bimbingan dan sokongan yang diberikan sepanjang saya berkhidmat "
            "di sini. Saya mendoakan agar {company} terus maju dan berjaya, dan berharap dapat "
            "terus berhubung pada masa hadapan."
        ),
        "signoff": "Yang benar,",
        # --- email drafts: header labels, subject, and the cover-email body ---
        "email_label_to": "Kepada",
        "email_label_subject": "Perkara",
        "email_subject": "Notis Peletakan Jawatan - {your_name} ({position})",
        "email_cover_notice": (
            "Dengan segala hormatnya, saya ingin memaklumkan secara rasmi peletakan jawatan "
            "saya sebagai {position} di {company}. Surat peletakan jawatan rasmi saya "
            "disertakan bersama e-mel ini untuk simpanan pihak tuan/puan. Selaras dengan "
            "tempoh notis saya selama {period}, hari terakhir saya bekerja ialah {last_day}."
        ),
        "email_cover_immediate": (
            "Dengan segala hormatnya, saya ingin memaklumkan secara rasmi peletakan jawatan "
            "saya sebagai {position} di {company}, berkuat kuasa serta-merta. Surat peletakan "
            "jawatan rasmi saya disertakan bersama e-mel ini untuk simpanan pihak tuan/puan, "
            "dan hari terakhir saya bekerja ialah {last_day}."
        ),
        "email_cover_body": (
            "Saya komited untuk memastikan proses peralihan tugas berjalan lancar sebelum saya "
            "berhenti, dan sedia membantu dengan apa-apa cara bagi melicinkan proses peralihan "
            "ini. Terima kasih atas sokongan dan bimbingan yang saya terima sepanjang tempoh "
            "saya berkhidmat di {company}."
        ),
    },
    "zh": {
        "filename_prefix": "辞职信",
        "title": "辞 职 信",
        "subject": "事由：辞去{position}职务",
        "salutation_named": "尊敬的{name}：",
        "salutation_generic": "尊敬的领导：",
        "opening_notice": (
            "兹正式通知贵公司，本人决定辞去在{company}担任的{position}职务。根据本人雇佣合约的"
            "规定，本人将履行{period}的通知期，最后工作日为{last_day}。"
        ),
        "opening_immediate": (
            "兹正式通知贵公司，本人决定即时辞去在{company}担任的{position}职务，最后工作日为"
            "{last_day}。"
        ),
        "gratitude_ach": (
            "在{company}工作的这段时间令本人获益良多。本人尤其感激能有机会{achievements}，"
            "亦十分珍惜在职期间所获得的专业成长、悉心指导与团队协作。"
        ),
        "gratitude_noach": (
            "在{company}工作的这段时间令本人获益良多，本人十分珍惜在职期间所获得的专业成长、"
            "悉心指导与团队协作。"
        ),
        "handover": (
            "在余下的任期内，本人承诺确保工作顺利、专业地交接。本人乐意协助移交各项职责、"
            "整理在办事务的记录、培训接任人员，并尽力完成一切未尽事宜。"
        ),
        "closing_para": (
            "感谢您在本人任职期间给予的指导与支持。谨祝您与{company}事业蒸蒸日上，"
            "并期盼日后仍能保持联系。"
        ),
        "signoff": "此致敬礼！",
        # --- email drafts: header labels, subject, and the cover-email body ---
        "email_label_to": "收件人",
        "email_label_subject": "主题",
        "email_subject": "辞职通知：{your_name}（{position}）",
        "email_cover_notice": (
            "本人谨此正式通知，本人决定辞去在{company}担任的{position}职务。本人的正式辞职信"
            "已随本邮件附上，敬请查收存档。根据本人{period}的通知期，本人的最后工作日为"
            "{last_day}。"
        ),
        "email_cover_immediate": (
            "本人谨此正式通知，本人决定即时辞去在{company}担任的{position}职务。本人的正式"
            "辞职信已随本邮件附上，敬请查收存档，本人的最后工作日为{last_day}。"
        ),
        "email_cover_body": (
            "在离职前，本人承诺确保各项工作顺利交接，并乐意尽力协助，使交接过程尽可能顺畅。"
            "感谢本人在{company}任职期间所获得的支持与指导。"
        ),
    },
}


# --------------------------------------------------------------------------- #
# Notice-period wording, per language and unit
# --------------------------------------------------------------------------- #
# English needs singular/plural; Malay and Chinese use one invariant word.
PERIOD_WORDS = {
    "en": {"day": ("day", "days"), "week": ("week", "weeks"), "month": ("month", "months")},
    "ms": {"day": "hari", "week": "minggu", "month": "bulan"},
    "zh": {"day": "天", "week": "周", "month": "个月"},
}


def format_period(count, unit, lang):
    """Render the notice period for a language, e.g. '2 weeks', '2 minggu', '2周'."""
    words = PERIOD_WORDS[lang][unit]
    if lang == "en":
        singular, plural = words
        return f"{count} {singular if count == 1 else plural}"
    if lang == "zh":
        return f"{count}{words}"  # no space before Chinese measure words
    return f"{count} {words}"


# --------------------------------------------------------------------------- #
# Console / encoding setup (so Chinese never crashes a legacy Windows console)
# --------------------------------------------------------------------------- #
def configure_console():
    for name in ("stdout", "stderr", "stdin"):
        stream = getattr(sys, name, None)
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass


# --------------------------------------------------------------------------- #
# Input helpers
# --------------------------------------------------------------------------- #
def _strip_bom(text):
    """Remove a leading byte-order mark, including a UTF-8 BOM that was decoded
    under a legacy Windows code page (which appears as the characters 'ï»¿')."""
    for bom in ("﻿", "ï»¿"):  # UTF-8 BOM, and its cp1252 mis-decode
        if text.startswith(bom):
            return text[len(bom):]
    return text


def ask(label, default="", required=False):
    """Prompt for a line of text. Enter accepts the [default] shown."""
    suffix = f" [{default}]" if default else ""
    while True:
        try:
            value = _strip_bom(input(f"{label}{suffix}: ")).strip()
        except EOFError:
            value = ""
        if not value:
            value = default
        if required and not value:
            print("  -> This field is required. Please enter a value.")
            continue
        return value


def ask_date(label, default_today=False):
    """Prompt for a date in yyyy-MM-dd format and return a date object."""
    default = date.today().isoformat() if default_today else ""
    while True:
        raw = ask(label, default=default)
        try:
            return datetime.strptime(raw, "%Y-%m-%d").date()
        except ValueError:
            print("  -> Invalid date. Use the format yyyy-MM-dd, e.g. 2026-06-28.")


def ask_int(label, default=None, min_value=0):
    """Prompt for a whole number >= min_value."""
    shown = "" if default is None else str(default)
    while True:
        raw = ask(label, default=shown)
        try:
            value = int(raw)
        except ValueError:
            print("  -> Please enter a whole number, e.g. 1, 2 or 3.")
            continue
        if value < min_value:
            print(f"  -> Please enter a number of {min_value} or more.")
            continue
        return value


def ask_yes_no(label, default_yes=True):
    hint = "Y/n" if default_yes else "y/N"
    try:
        raw = _strip_bom(input(f"{label} [{hint}]: ")).strip().lower()
    except EOFError:
        return default_yes
    if not raw:
        return default_yes
    return raw in ("y", "yes")


def ask_language():
    print("Select letter language / Pilih bahasa surat / 选择信件语言:")
    print("  1. English (default)")
    print("  2. Bahasa Melayu")
    print("  3. 中文 (简体)")
    while True:
        try:
            raw = _strip_bom(input("Enter 1, 2 or 3 [1]: ")).strip()
        except EOFError:
            raw = ""
        if raw in LANGUAGE_CHOICES:
            return LANGUAGE_CHOICES[raw]
        print("  -> Please enter 1, 2 or 3.")


def ask_notice_unit():
    print("Is your notice period counted in days, weeks or months?")
    print("  1. Day(s)")
    print("  2. Week(s)")
    print("  3. Month(s) (default)")
    while True:
        try:
            raw = _strip_bom(input("Enter 1, 2 or 3 [3]: ")).strip()
        except EOFError:
            raw = ""
        if raw in NOTICE_UNIT_CHOICES:
            return NOTICE_UNIT_CHOICES[raw]
        print("  -> Please enter 1, 2 or 3.")


# --------------------------------------------------------------------------- #
# Date maths
# --------------------------------------------------------------------------- #
def add_months(start, months):
    """Add a number of whole months to a date, clamping to month length."""
    total = start.month - 1 + months
    year = start.year + total // 12
    month = total % 12 + 1
    last_day_of_month = calendar.monthrange(year, month)[1]
    return date(year, month, min(start.day, last_day_of_month))


def compute_last_working_day(start, count, unit):
    """Last working day = serve through the notice period, i.e. the day before it ends."""
    if count <= 0:
        return start
    if unit == "month":
        end = add_months(start, count)
    elif unit == "week":
        end = start + timedelta(weeks=count)
    else:  # day
        end = start + timedelta(days=count)
    return end - timedelta(days=1)


# --------------------------------------------------------------------------- #
# Fonts (for the PDF)
# --------------------------------------------------------------------------- #
class CJKFontError(Exception):
    """Raised when a Chinese letter is requested but no CJK font is available."""


def find_cjk_font():
    """Return a path to a system font that can render Chinese, or None."""
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",      # Microsoft YaHei (Windows)
        r"C:\Windows\Fonts\simsun.ttc",    # SimSun (Windows)
        r"C:\Windows\Fonts\simhei.ttf",    # SimHei (Windows)
        r"C:\Windows\Fonts\Deng.ttf",      # DengXian (Windows)
        "/System/Library/Fonts/PingFang.ttc",                 # macOS
        "/System/Library/Fonts/STHeiti Light.ttc",            # macOS
        "/System/Library/Fonts/Supplemental/Songti.ttc",      # macOS
        "/Library/Fonts/Arial Unicode.ttf",                   # macOS (older)
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",   # Linux
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",   # Linux
        "/usr/share/fonts/opentype/noto/NotoSansSC-Regular.otf",    # Linux
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


# --------------------------------------------------------------------------- #
# Letter content (written to an experienced-HR standard, per language)
# --------------------------------------------------------------------------- #
def build_letter(d, lang):
    t = TRANSLATIONS[lang]
    count = d["notice_count"]
    unit = d["notice_unit"]
    kw = dict(
        position=d["your_position"],
        company=d["company"],
        name=d.get("recipient_name", ""),
        period=format_period(count, unit, lang),
        last_day=d["last_day"],
        achievements=d.get("achievements", ""),
    )

    opening = (t["opening_immediate"] if count <= 0 else t["opening_notice"]).format(**kw)
    gratitude = (t["gratitude_ach"] if d.get("achievements") else t["gratitude_noach"]).format(**kw)
    handover = t["handover"].format(**kw)
    closing_para = t["closing_para"].format(**kw)
    salutation = t["salutation_named"].format(**kw) if d.get("recipient_name") else t["salutation_generic"]
    subject = t["subject"].format(**kw)

    sender_lines = [x for x in (
        d["your_name"], d["your_position"], d.get("your_dept", ""),
        d.get("your_email", ""), d.get("your_phone", ""),
    ) if x]

    recipient_lines = [x for x in (
        d.get("recipient_name", ""), d.get("recipient_position", ""),
        d["company"], d.get("addr1", ""), d.get("city", ""),
    ) if x]

    return {
        "title": t["title"],
        "sender_lines": sender_lines,
        "date": d["letter_date"],
        "recipient_lines": recipient_lines,
        "subject": subject,
        "salutation": salutation,
        "paragraphs": [opening, gratitude, handover, closing_para],
        "closing": t["signoff"],
        "name": d["your_name"],
        "position": d["your_position"],
    }


# --------------------------------------------------------------------------- #
# Email content (same language as the letter)
#   - "with_attachment": a short cover email to send WITH the letter attached
#   - "without_attachment": the full resignation carried in the email body,
#     reusing the very paragraphs of the letter so the two never drift apart
# --------------------------------------------------------------------------- #
def build_emails(d, lang):
    t = TRANSLATIONS[lang]
    count = d["notice_count"]
    unit = d["notice_unit"]
    kw = dict(
        position=d["your_position"],
        company=d["company"],
        your_name=d["your_name"],
        name=d.get("recipient_name", ""),
        period=format_period(count, unit, lang),
        last_day=d["last_day"],
        achievements=d.get("achievements", ""),
    )

    salutation = t["salutation_named"].format(**kw) if d.get("recipient_name") else t["salutation_generic"]

    # "To:" line, shown only when we actually have an address to write.
    to_line = ""
    recipient_email = d.get("recipient_email", "")
    if recipient_email:
        rname = d.get("recipient_name", "")
        to_line = f"{rname} <{recipient_email}>" if rname else recipient_email

    # Signature: sign-off, then name / position / department / contact details.
    signature = [t["signoff"], d["your_name"], d["your_position"]]
    for extra in (d.get("your_dept", ""), d.get("your_email", ""), d.get("your_phone", "")):
        if extra:
            signature.append(extra)

    # Email 1 - cover note to accompany the attached letter.
    cover_open = (t["email_cover_immediate"] if count <= 0 else t["email_cover_notice"]).format(**kw)
    with_attachment = [cover_open, t["email_cover_body"].format(**kw)]

    # Email 2 - the full resignation written straight into the email body.
    opening = (t["opening_immediate"] if count <= 0 else t["opening_notice"]).format(**kw)
    gratitude = (t["gratitude_ach"] if d.get("achievements") else t["gratitude_noach"]).format(**kw)
    without_attachment = [opening, gratitude, t["handover"].format(**kw), t["closing_para"].format(**kw)]

    return {
        "label_to": t["email_label_to"],
        "label_subject": t["email_label_subject"],
        "to": to_line,
        "subject": t["email_subject"].format(**kw),
        "salutation": salutation,
        "signature": signature,
        "with_attachment": with_attachment,
        "without_attachment": without_attachment,
    }


def render_email(em, which):
    """Render one email variant ('with_attachment' / 'without_attachment') as
    plain text: header line(s), a blank line, salutation, body and signature."""
    lines = []
    if em["to"]:
        lines.append(f"{em['label_to']}: {em['to']}")
    lines.append(f"{em['label_subject']}: {em['subject']}")
    lines.append("")
    lines.append(em["salutation"])
    lines.append("")
    for para in em[which]:
        lines.append(para)
        lines.append("")
    signoff, *details = em["signature"]
    lines.append(signoff)
    lines.append("")
    lines.extend(details)
    return "\n".join(lines).rstrip() + "\n"


# --------------------------------------------------------------------------- #
# Renderers
# --------------------------------------------------------------------------- #
def write_docx(c, path, lang):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    if lang == "zh":
        normal.font.name = "Microsoft YaHei"
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    else:
        normal.font.name = "Calibri"

    def para(text="", bold=False, align=None, size=None, space_after=4):
        p = doc.add_paragraph()
        if text:
            run = p.add_run(text)
            run.bold = bold
            if size:
                run.font.size = Pt(size)
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    para(c["title"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=12)
    for i, line in enumerate(c["sender_lines"]):
        para(line, bold=(i == 0), space_after=0)
    para("", space_after=4)
    para(c["date"], space_after=8)
    for line in c["recipient_lines"]:
        para(line, space_after=0)
    para("", space_after=4)
    para(c["subject"], bold=True, space_after=8)
    para(c["salutation"], space_after=8)
    for body in c["paragraphs"]:
        para(body, space_after=8)
    para(c["closing"], space_after=0)
    para("", space_after=0)
    para("", space_after=0)
    para(c["name"], bold=True, space_after=0)
    para(c["position"], space_after=0)

    doc.save(str(path))


def pdf_safe(text):
    """Make text safe for the PDF core font (Latin-1); normalise smart punctuation."""
    if not text:
        return ""
    replacements = {
        "‘": "'", "’": "'", "“": '"', "”": '"',
        "–": "-", "—": "-", "…": "...", " ": " ",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "replace").decode("latin-1")


def write_pdf(c, path, lang):
    core = (lang != "zh")
    font_path = None
    if not core:
        font_path = find_cjk_font()
        if not font_path:
            raise CJKFontError()

    pdf = FPDF(format="A4", unit="mm")
    pdf.set_auto_page_break(True, margin=25)
    pdf.set_margins(25, 25, 25)
    pdf.add_page()

    if core:
        family = "Helvetica"
    else:
        family = "Letter"
        pdf.add_font(family, "", font_path)

    def block(text, size=11, bold=False, align="L", height=6, gap=0):
        style = "B" if (bold and core) else ""
        if not core and align == "J":
            align = "L"
        pdf.set_font(family, style, size)
        out = pdf_safe(text) if core else text
        pdf.multi_cell(0, height, out, align=align, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if gap:
            pdf.ln(gap)

    block(c["title"], size=14, bold=True, align="C", height=8, gap=4)
    for i, line in enumerate(c["sender_lines"]):
        block(line, bold=(i == 0))
    pdf.ln(3)
    block(c["date"], gap=3)
    for line in c["recipient_lines"]:
        block(line)
    pdf.ln(3)
    block(c["subject"], bold=True, gap=3)
    block(c["salutation"], gap=2)
    for body in c["paragraphs"]:
        block(body, align="J", gap=3)
    block(c["closing"])
    pdf.ln(14)
    block(c["name"], bold=True)
    block(c["position"])

    pdf.output(str(path))


# --------------------------------------------------------------------------- #
# Output helpers
# --------------------------------------------------------------------------- #
def safe_filename(name):
    """Keep letters (incl. Chinese), digits and spaces; drop invalid filename chars."""
    cleaned = re.sub(r'[<>:"/\\|?*]', "", name)        # characters illegal in filenames
    cleaned = re.sub(r"[\x00-\x1f]", "", cleaned)       # control characters
    cleaned = re.sub(r"\s+", "_", cleaned.strip())      # whitespace -> underscore
    cleaned = cleaned.strip("._-")
    return cleaned or "letter"


def unique_base(folder, base):
    candidate = base
    counter = 2
    while (folder / f"{candidate}.docx").exists() or (folder / f"{candidate}.pdf").exists():
        candidate = f"{base}_{counter}"
        counter += 1
    return candidate


def write_text(path, text):
    """Write UTF-8 text (used for the email drafts) with normal newlines."""
    path.write_text(text, encoding="utf-8")


def open_folder(folder):
    try:
        if sys.platform.startswith("win"):
            os.startfile(str(folder))  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.run(["open", str(folder)], check=False)
    except Exception:
        pass


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def main():
    configure_console()

    print("=" * 50)
    print("   HR Letter Generator  -  Resignation Letter")
    print("=" * 50)
    lang = ask_language()
    print()
    print("Answer the prompts below. Press Enter to accept a [default].")
    print("Optional fields can be left blank by pressing Enter.\n")

    print("-- Your details --")
    your_name = ask("Your full name", required=True)
    your_position = ask("Your position / job title", required=True)
    your_dept = ask("Your department (optional)")
    your_email = ask("Your email (optional)")
    your_phone = ask("Your phone (optional)")

    print("\n-- Recipient details --")
    recipient_name = ask("Recipient's full name, e.g. your manager (optional)")
    recipient_position = ask("Recipient's position / title (optional)")
    recipient_email = ask("Recipient's email, for the email 'To:' line (optional)")
    company = ask("Company / organisation name", required=True)
    addr1 = ask("Company address line 1 (optional)")
    city = ask("City, Postcode (optional)")

    print("\n-- Notice period --")
    letter_date = ask_date("Date you are giving notice (yyyy-MM-dd)", default_today=True)
    notice_unit = ask_notice_unit()
    notice_count = ask_int(f"Notice period in {UNIT_PLURAL_EN[notice_unit]}", default=1, min_value=0)
    last_day = compute_last_working_day(letter_date, notice_count, notice_unit)
    print(f"\n  Computed last working day: {last_day.isoformat()}")
    if not ask_yes_no("  Use this last working day?", default_yes=True):
        last_day = ask_date("  Enter the correct last working day (yyyy-MM-dd)")

    print("\n-- Optional --")
    achievements = ask(
        "Key achievement(s) to mention, in the chosen language,\n"
        "  e.g. 'lead the payroll system migration' (optional)"
    )

    data = {
        "your_name": your_name,
        "your_position": your_position,
        "your_dept": your_dept,
        "your_email": your_email,
        "your_phone": your_phone,
        "recipient_name": recipient_name,
        "recipient_position": recipient_position,
        "recipient_email": recipient_email,
        "company": company,
        "addr1": addr1,
        "city": city,
        "letter_date": letter_date.isoformat(),
        "notice_count": notice_count,
        "notice_unit": notice_unit,
        "last_day": last_day.isoformat(),
        "achievements": achievements,
    }

    print("\n" + "=" * 50)
    print("  Review")
    print("=" * 50)
    print(f"  Language       : {lang}")
    print(f"  Name           : {your_name}")
    print(f"  Position       : {your_position}")
    print(f"  Company        : {company}")
    print(f"  Letter date    : {data['letter_date']}")
    print(f"  Notice period  : {format_period(notice_count, notice_unit, 'en')}")
    print(f"  Last working   : {data['last_day']}")
    print("=" * 50)
    if not ask_yes_no("\nGenerate the letter now?", default_yes=True):
        print("Cancelled. No files were created.")
        return

    content = build_letter(data, lang)
    outdir = Path(__file__).resolve().parent / "generated_letters"
    outdir.mkdir(exist_ok=True)
    prefix = TRANSLATIONS[lang]["filename_prefix"]
    base = unique_base(outdir, f"{prefix}_{safe_filename(your_name)}")
    docx_path = outdir / f"{base}.docx"
    pdf_path = outdir / f"{base}.pdf"

    write_docx(content, docx_path, lang)
    print("\nDone! Your files have been created:")
    print(f"  - {docx_path}")
    try:
        write_pdf(content, pdf_path, lang)
        print(f"  - {pdf_path}")
    except CJKFontError:
        print("  - (PDF skipped: no Chinese-capable font was found on this computer.)")
        print("    The Word (.docx) file above contains the complete Chinese letter.")

    # Email drafts: one to send WITH the letter attached, one that carries the
    # full resignation in the body (no attachment). Saved as .txt and shown below.
    emails = build_emails(data, lang)
    cover_text = render_email(emails, "with_attachment")
    full_text = render_email(emails, "without_attachment")
    cover_path = outdir / f"{base}_Email_1_with_letter_attached.txt"
    full_path = outdir / f"{base}_Email_2_full_text_no_attachment.txt"
    write_text(cover_path, cover_text)
    write_text(full_path, full_text)
    print(f"  - {cover_path}")
    print(f"  - {full_path}")

    print("\n" + "=" * 60)
    print("  EMAIL 1  -  send this WITH your resignation letter attached")
    print("  (attach the .docx or .pdf above before sending)")
    print("=" * 60 + "\n")
    print(cover_text)
    print("=" * 60)
    print("  EMAIL 2  -  send this on its own, no attachment needed")
    print("  (the full resignation is written in the email body)")
    print("=" * 60 + "\n")
    print(full_text)

    open_folder(outdir)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nCancelled by user.")
        sys.exit(1)
